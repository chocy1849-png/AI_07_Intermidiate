from __future__ import annotations

from pathlib import Path

try:
    from docx import Document
except ImportError:  # pragma: no cover - dependency handled in runtime
    Document = None


class DOCXParseError(RuntimeError):
    """Raised when a DOCX document cannot be parsed."""


def parse_docx(file_path: str | Path) -> dict:
    if Document is None:
        raise DOCXParseError("python-docx is not installed. Run `pip install -r requirements.txt` first.")

    path = Path(file_path)
    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    tables = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                tables.append(" | ".join(cells))

    text = "\n".join(paragraphs + tables).strip()
    if not text:
        raise DOCXParseError(f"No readable text extracted from {path.name}")

    return {
        "text": text,
        "metadata": {
            "source_path": str(path.resolve()),
            "file_name": path.name,
            "file_format": "docx",
            "parser": "python-docx",
            "page_count": None,
            "char_count": len(text),
        },
    }
